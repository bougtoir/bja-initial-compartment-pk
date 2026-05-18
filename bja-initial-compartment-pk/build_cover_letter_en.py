#!/usr/bin/env python3
"""Generate English BJA cover letter as .docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import date

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

def add_para(text, bold=False, italic=False, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ===== DATE =====
add_para('[Date]', space_after=12)

# ===== ADDRESSEE =====
add_para('Professor Hugh Hemmings')
add_para('Editor-in-Chief')
add_para('British Journal of Anaesthesia')
add_para('Weill Cornell Medical College')
add_para('New York, NY, USA', space_after=12)

# ===== SUBJECT =====
add_para('Dear Professor Hemmings,', space_after=12)

add_para(
    'Re: Submission of Narrative Review \u2014 '
    '\u201cRethinking Maximum Dose Limits for Local Anaesthetics in Regional Anaesthesia: '
    'The Case for Initial Compartment-Dependent Pharmacokinetic Modelling and a Call for '
    'Route-Adaptive PKPD Simulation in Anaesthesia Information Management Systems\u201d',
    bold=True, space_after=12
)

# ===== BODY =====
add_para(
    'We are pleased to submit the above manuscript for consideration as a Narrative Review '
    'in the British Journal of Anaesthesia. This work addresses a fundamental but overlooked '
    'problem in regional anaesthesia pharmacokinetics: current maximum recommended doses for '
    'local anaesthetics are derived from models that assume intravenous administration, yet the '
    'initial pharmacokinetic compartment in regional anaesthesia depends critically on the route '
    'of administration and the success of the block.',
    space_after=12
)

add_para(
    'Our review proposes a conceptual framework in which the initial compartment of drug '
    'deposition\u2014vessel-poor tissue for blocks with slow systemic absorption, vessel-rich '
    'tissue or plasma for blocks with rapid systemic absorption\u2014determines the pharmacokinetic '
    'trajectory and, consequently, the risk of local anaesthetic systemic toxicity (LAST). '
    'We demonstrate that this framework can be implemented using freely available physiologically '
    'based pharmacokinetic (PBPK) modelling platforms such as PK-Sim and MoBi, and we argue '
    'that pharmacokinetic\u2013pharmacodynamic (PKPD) simulation modules embedded in modern '
    'anaesthesia information management systems (AIMS) should incorporate route-of-administration '
    'logic for real-time dose guidance.',
    space_after=12
)

add_para(
    'This manuscript builds upon and extends the concepts introduced in the recent BJA editorial '
    'by De Cassai et al. (2025), which highlighted the need for population pharmacokinetic studies '
    'of local anaesthetics after fascial plane blocks. We believe our work provides a timely and '
    'comprehensive framework that will be of significant interest to BJA readers, including '
    'anaesthetists, pharmacologists, and developers of clinical decision support systems.',
    space_after=12
)

add_para(
    'The manuscript comprises approximately 4800 words in the main text, with 50 references, '
    '4 figures (including 2 simulation-derived figures and 2 conceptual flow diagrams), and '
    '1 table. These are within the BJA limits for Narrative Reviews (5000 words, 150 references, '
    '6 tables/figures). All figures are provided in colour at 300 dpi or higher resolution.',
    space_after=12
)

add_para(
    'We confirm that this manuscript is original work, has not been published previously, and '
    'is not under consideration for publication elsewhere. All authors have read and approved '
    'the final manuscript and agree to its submission to the British Journal of Anaesthesia.',
    space_after=12
)

add_para(
    'We declare no conflicts of interest related to this work. '
    '[If applicable, please modify this statement to disclose any relevant conflicts.]',
    space_after=12
)

add_para(
    'We would welcome the opportunity to revise the manuscript in response to reviewer feedback. '
    'Thank you for considering our submission.',
    space_after=12
)

add_para('Yours sincerely,', space_after=24)

add_para('[Corresponding author name]')
add_para('[Title, Department]')
add_para('[Institution]')
add_para('[Address]')
add_para('[Email]')
add_para('[Telephone]')

# ===== SAVE =====
out = os.path.join(os.path.dirname(__file__), 'BJA_Cover_Letter_English.docx')
doc.save(out)
print(f'English cover letter saved to {out}')
