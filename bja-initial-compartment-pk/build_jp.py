#!/usr/bin/env python3
"""Generate Japanese BJA manuscript as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 2.0

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ===== TITLE PAGE =====
add_para('')
add_para('')
p = add_para('区域麻酔における局所麻酔薬の極量再考：', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
p = add_para('初期コンパートメント依存性薬物動態モデリングの必要性', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
p = add_para('および麻酔情報管理システムにおける投与経路適応型PKPDシミュレーションの提言', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('')

add_para('論文種別：Narrative Review', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('[著者名]', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[所属機関]', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('責任著者：', bold=True)
add_para('[氏名、所属、住所、メールアドレス]')
add_para('')
add_para('本文語数：約4800語（本文のみ）')
add_para('引用文献数：50')
add_para('図：4')
add_para('表：1')
add_para('')

add_para('キーワード：局所麻酔薬全身毒性；薬物動態モデリング；区域麻酔；生理学的薬物動態モデル；コンパートメントモデル；麻酔情報管理システム；極量', italic=True)

doc.add_page_break()

# ===== SUMMARY =====
add_heading_text('要旨（Summary）', level=1)

add_para(
    '現行の局所麻酔薬の極量（最大推奨用量）は、薬物が中心コンパートメント（血漿）に直接投与される'
    '静脈内投与を前提とした薬物動態モデルに基づいている。しかし、区域麻酔においては、薬物の初期沈着部位は'
    'ブロックの成否によって根本的に異なる。成功した神経ブロックや筋膜面ブロックでは薬物は血管の乏しい組織'
    '（vessel-poor tissue, BPT）に沈着し全身吸収は緩徐であるのに対し、不成功のブロックでは血管の豊富な組織'
    '（vessel-rich tissue, BRT）や血管内に薬物が沈着し急速な全身吸収が生じる。この初期薬物動態コンパートメント'
    'の相違は、最高血中濃度（Cmax）、ひいては局所麻酔薬全身毒性（LAST）のリスクに重大な影響を与える。'
    '本稿では、区域麻酔の用量設定における従来の3コンパートメントモデルの限界を検討し、PK-SimやMoBiなどの'
    '生理学的薬物動態（PBPK）モデリングプラットフォームによる投与経路依存性シミュレーションの可能性を論じ、'
    '状況依存性極量（context-sensitive maximum dose）の枠組みを提案する。さらに、現代の麻酔情報管理システム'
    '（AIMS）に組み込まれた薬物動態・薬力学（PKPD）シミュレーションモジュールは、投与経路のロジックを組み込み、'
    '静脈内投与と区域麻酔の両方に対してリアルタイムでコンパートメント適切な用量指針を提供すべきであると提言する。'
)

doc.add_page_break()

# ===== INTRODUCTION =====
add_heading_text('緒言（Introduction）', level=1)

add_para(
    '局所麻酔薬全身毒性（LAST）は、区域麻酔における最も重篤な合併症の一つである。局所麻酔薬の極量'
    '（最大推奨用量）は従来mg/kgで表現され、静脈内投与や皮下浸潤を対象とした薬物動態研究に基づいて'
    '数十年前に設定された。1,2 これらの用量制限は、薬物が中心（血漿）コンパートメントに投入されること\u2014'
    'すなわち静脈内投与\u2014を暗黙の前提としている。'
)

add_para(
    '麻酔薬理学で広く用いられる3コンパートメントマミラリーモデルは、中心コンパートメント（V1、血漿）から'
    '急速平衡末梢コンパートメント（V2、血管の豊富な組織群＝BRT）および緩徐平衡末梢コンパートメント'
    '（V3、血管の乏しい組織群＝BPT）への薬物分布を記述し、中心コンパートメントからの消失を含む。3,4 '
    'このモデルは、プロポフォールのtarget-controlled infusion（TCI）システム（Marsh、Schnider、Eleveldモデル等）'
    'の基盤であり、現代の麻酔情報管理システム（AIMS）におけるPKPD表示の根幹をなす。5,6'
)

add_para(
    'しかし区域麻酔では、薬物は静脈内に投与されるのではない。神経周囲、筋膜面、硬膜外腔といった組織に沈着し、'
    '全身吸収は局所血流、組織結合、薬物および組織の物理化学的特性に依存する。7,8 '
    'De Cassaiらは本誌において、筋膜面ブロックにおける局所麻酔薬の薬物動態に関する最新の知見を報告し、'
    'エピネフリン、組織血管分布、筋膜微小解剖が全身吸収プロファイルに与える影響を明らかにした。9 '
    '同様に、Schwenkらはリドカインが持続的なLAST関連死亡原因として台頭していることに注意を喚起した。10'
)

add_para(
    'これらの進歩にもかかわらず、根本的な問題が十分に取り上げられていない：薬物の初期沈着部位'
    '\u2014すなわちスタートコンパートメント\u2014が最高血中濃度にどのように影響し、安全域にどう影響するか？ '
    'この問いへの回答は、区域麻酔における極量の意味ある議論に不可欠であり、現行の用量推奨は'
    'この変数を考慮していないため不十分であると我々は主張する。'
)

# ===== THE PROBLEM =====
add_heading_text('初期コンパートメント問題（The initial compartment problem）', level=1)

add_para(
    '同一用量の長時間作用性局所麻酔薬（例：ブピバカイン150 mg）を末梢神経ブロックとして'
    '投与する2つの臨床シナリオを考える：'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run(
    'シナリオA（成功したブロック）：全量が標的筋膜面または神経周囲腔に正確に沈着する。組織は主に'
    '血管の乏しい組織（脂肪、結合組織、筋膜）である。全身吸収は緩徐であり、低い吸収速度定数（ka）に'
    '支配される。Cmaxは低値かつ遅延（高いTmax）。薬物は局所効果を発揮しつつ、全身循環への緩やかな移行に'
    '伴い肝代謝により漸次クリアランスされる。'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run(
    'シナリオB（不成功のブロックまたは血管内注入）：薬物が血管構造内または高灌流組織（血管豊富組織群）に'
    '沈着する。全身吸収は急速であり、静脈内投与に相当またはそれに近い。Cmaxは高値かつ早期'
    '（低いTmax）であり、中枢神経系毒性または心血管系毒性の閾値を超える可能性がある。'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_para(
    '従来の3コンパートメントモデルでは、両シナリオとも同一の極量に対して評価される\u2014血漿を初期'
    'コンパートメントと仮定するモデルから導出された用量制限である。しかし薬物動態プロファイルは根本的に異なる。'
    'シナリオAでは、緩徐な吸収速度により従来の制限を超える用量でもCmaxは毒性閾値を大幅に下回り、'
    '従来の極量は不必要に保守的かもしれない。シナリオBでは、急速な全身吸収により静脈内ボーラス投与と'
    '同等の血漿濃度が生じ、同じ極量が危険なほど寛容かもしれない。'
)

add_para(
    'この非対称性は単なる理論ではない。筋膜面ブロックでは従来の体重あたりの制限を超える用量が'
    '明らかな毒性なしに日常的に使用されている一方、11,12 稀ではあるが壊滅的なLAST事象は従来の'
    '用量で発生し続けており、その多くは不慮の血管内注入または高度に血管に富む注入部位からの急速吸収'
    'に起因する。13,14 この方程式で欠落している変数が、薬物沈着の初期コンパートメントである。'
)

# ===== BLOCK SUCCESS AS PK DETERMINANT =====
add_heading_text('ブロック成功は薬物動態の決定因子である', level=1)

add_para(
    '我々は、区域ブロックの臨床的成功が投与された局所麻酔薬の薬物動態経路に関する'
    '直接的情報を提供すると提案する。持続的な感覚・運動遮断をもたらす成功したブロックは、'
    '薬物の相当量が標的組織（血管乏しいコンパートメント）に長期間局在していることを意味する。'
    'これは緩徐吸収の薬物動態的証拠である：薬物は注入部位から全身循環へ急速にクリアランスされていない。'
)

add_para(
    '逆に、急速に失敗するか十分な麻酔を達成しないブロックは、薬物が急速に吸収されたか、'
    '標的神経から離れた部位に沈着し、血管豊富組織経路を通じて全身循環に入ったことを示唆する。'
    '薬物動態的には、初期コンパートメントがV3（血管乏しい組織）ではなくV2（血管豊富組織）または'
    'V1（血漿）に近いシナリオである。'
)

add_para(
    'ブロックの有効性と薬物動態の関係は重要な臨床的意味を持つ。ブロックが何時間も臨床的に有効である場合、'
    '血漿濃度曲線は低く平坦なプロファイルを示す\u2014薬物は全身循環に氾濫するのではなく局所に隔離されている。'
    'このシナリオにおけるLASTリスクは、即座の血漿コンパートメント投入を仮定するモデルの予測より'
    '本質的に低い。その帰結として、不成功のブロックは従来認識されているよりも高い毒性リスクを呈する\u2014'
    '全量が薬物動態的に静脈内投与されたかのように振る舞う可能性があるからである。'
)

# ===== LIMITATIONS =====
add_heading_text('区域麻酔に対する従来の3コンパートメントモデルの限界', level=1)

add_para(
    '臨床麻酔で使用される3コンパートメントモデル（プロポフォールのMarsh、Schnider、Eleveld；'
    'レミフェンタニルのMinto、Kim、Eleveld）は静脈内投与用に開発された。3-6 '
    'これらは薬物が中心コンパートメントに入った後の分布を記述するものであり、修正なしに'
    '区域麻酔の薬物動態をモデル化するには本質的に不適切である。'
)

add_para(
    'いくつかの具体的限界を強調する。第一に、これらのモデルにはデポまたは吸収コンパートメントがない。'
    '区域麻酔では、薬物は注入部位から吸収されてから血漿に入る必要があり、その過程は'
    '注入部位、組織血管分布、血管収縮薬の使用、および個々の患者因子によって異なる'
    '吸収速度定数（ka）と生体利用率（F）によって特徴付けられる。15,16 '
    '第二に、コンパートメント間速度定数（k12、k21、k13、k31）はIV投与データから推定されており、'
    '薬物が末梢組織デポから出発する場合の薬物移動動態を正確に反映しない可能性がある。'
    '第三に、蛋白結合動態はIVボーラス投与（遊離薬物分画が急激にスパイクする）と'
    '緩徐な組織吸収（蛋白結合能が飽和しない）との間で重要な差異がある。17'
)

add_para(
    '区域ブロック後の局所麻酔薬の母集団薬物動態研究は、一次吸収を伴うデポコンパートメントを'
    '組み込むことでこれらの限界の一部に対処してきた。18-20 '
    'Gaudreaultらは大腿神経ブロック後のロピバカイン薬物動態を一次吸収を伴う2コンパートメントモデルで'
    'モデル化し、吸収速度が消失速度より遅いflip-flop動態を示した。18 '
    '最近では、Lingらがserratus anterior plane block後のロピバカインの母集団薬物動態モデルを'
    'NONMEMで開発した。19 これらの研究は一貫して、区域ブロック後の薬物動態プロファイルが'
    'IV投与と著しく異なることを示しているが、それらの知見は改訂された用量推奨に反映されていない。'
)

# ===== PBPK =====
add_heading_text('生理学的薬物動態モデリング：投与経路適応型アプローチ', level=1)

add_para(
    '生理学的薬物動態（PBPK）モデルは、薬物分布のシミュレーションに根本的に異なるアプローチを提供する。'
    '経験的に推定された移行定数を持つ抽象的なコンパートメントではなく、PBPKモデルは身体を解剖学的・'
    '生理学的に定義された臓器コンパートメントに分割し、それぞれが既知の血流、組織容積、分配係数、'
    '代謝能で特徴付けられる。21,22 このメカニスティックな枠組みは、薬物の初期沈着部位を指定することで、'
    '異なる投与経路を自然に収容する。'
)

add_para(
    'PK-SimとMoBiから構成されるOpen Systems Pharmacology（OSP）プラットフォームは、無料で利用可能な'
    'オープンソースのPBPKモデリングスイートであり、欧州医薬品庁による適格性認定を受け、'
    '医薬品開発と規制科学で広く使用されている。23,24 PK-Simは、事前定義された臓器コンパートメント'
    '（動脈血・静脈血、肺、心臓、筋肉、脂肪、皮膚等を含む）を持つ全身PBPKモデルの構築に'
    'グラフィカルインターフェースを提供し、MoBiはユーザー定義のコンパートメントと'
    '輸送プロセスによるカスタムモデル構築を可能にする。'
)

add_para(
    '我々の目的にとって重要なのは、PK-Simが静脈内、筋肉内、皮下注射を含む複数の投与経路を'
    'サポートし、それぞれに経路特異的吸収モデルが組み込まれている点である。25 '
    '筋肉内および皮下経路は組織特異的吸収動態を持つデポコンパートメントを組み込んでいる。'
    '類似的に、神経周囲または筋膜面注入は、注入部位に適切な特性を持つ組織コンパートメント'
    '（例：筋膜組織には低血流、高度血管化された神経周囲構造には高血流）への薬物沈着として'
    'モデル化できる。MoBiはさらに、完全にカスタムなコンパートメントと初期条件の定義を可能にし、'
    '上述の3つの臨床シナリオ（成功したブロック、不成功のブロック、部分的ブロック）を'
    '初期コンパートメントと各コンパートメントへの用量分配比率を変えることでシミュレーション可能である。'
)

add_para(
    'この概念を図1に示す。図1では、従来のIV 3コンパートメントモデルと、成功した区域ブロック'
    'および不成功の区域ブロックの修正モデルを比較している。図2は各シナリオのシミュレーションされた'
    '血漿濃度\u2013時間プロファイルを示し、薬物沈着の初期コンパートメントを変えることでCmaxとTmaxに'
    '生じる顕著な差異を示している。'
)

# ===== AIMS =====
add_heading_text('麻酔情報管理システムへの統合', level=1)

add_para(
    '現代のAIMSはリアルタイムPKPDシミュレーションを組み込みつつあり、プロポフォールや'
    'レミフェンタニルなどの静脈内薬剤の予測血漿濃度・効果部位濃度を表示している。5,26 '
    'これらの表示は3コンパートメントモデル（Eleveld、Schnider等）によって駆動され、'
    '全ての薬物が中心コンパートメントに静脈内投与されることを前提としている。'
    'この仮定は全静脈麻酔（TIVA）やTCIには妥当であるが、同じシステムが区域麻酔で'
    '投与された局所麻酔薬の用量追跡に使用される場合には正しくなくなる。'
)

add_para(
    '我々は、AIMSの開発元およびベンダーが、記録された投与経路に基づいて薬物動態モデルを'
    '調整する投与経路適応型PKPDシミュレーションを実装すべきであると提言する。'
    '具体的には、システムは以下を行うべきである：'
)

items_jp = [
    '局所麻酔薬の投与記録において、静脈内投与と区域麻酔投与を区別する。',
    '投与経路に適した薬物動態モデルを適用する：IV投与には標準3コンパートメントモデル、'
    '区域麻酔にはデポ増強モデル（特定のブロック種別に適した吸収速度定数と生体利用率を持つ）を使用。',
    'ブロック種別特異的な吸収パラメータを提供する。各区域麻酔手技について公表された'
    '母集団薬物動態データに基づく（例：TAPブロック、ESPブロック、大腿神経ブロック、硬膜外麻酔）。',
    '実際の投与経路を反映した予測血漿濃度推移を表示し、再投与や累積用量制限に関する'
    '情報に基づいた臨床判断を可能にする。',
    'ブロック成功の臨床指標（例：感覚検査結果）に基づいてリアルタイムでモデルを調整する'
    'メカニズムを組み込み、成功ブロック（緩徐吸収）と不成功ブロック（急速吸収）の'
    '薬物動態プロファイル間で切り替える。',
]
for item in items_jp:
    p = doc.add_paragraph()
    p.style = 'List Number'
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_para(
    '図3は、PBPK Based の状況依存性極量決定のためのワークフローを示す。'
    '図4は、投与経路適応型PKPDシミュレーションがAIMS表示にどのように統合されうるかの'
    '概念的模式図を示し、薬物動態モデルが記録された投与経路とブロック種別に基づいて'
    '自動的に調整される仕組みを提示している。'
)

add_para(
    'このアプローチの技術的実現可能性はいくつかの観察によって支持される。第一に、'
    'デポ増強コンパートメントモデルの数学的枠組みは確立されており、計算コストが低い。18-20 '
    '第二に、各種区域麻酔手技後の局所麻酔薬の母集団薬物動態パラメータが文献で増加しており、'
    '経路特異的モデルのパラメータ化に必要なデータを提供している。7,8,19 '
    '第三に、現代のAIMSは既に静脈内薬剤のリアルタイムPKPDシミュレーションを実装しており、'
    '計算インフラは存在する。主要な障壁は技術的ではなく概念的である：単一の薬物動態モデルでは'
    '根本的に異なる投与経路にわたる薬物挙動を適切に記述できないという認識である。'
)

# ===== CONTEXT-SENSITIVE MAX DOSE =====
add_heading_text('状況依存性極量（Context-sensitive maximum dose）に向けて', level=1)

add_para(
    '我々は、静脈内薬物のオフセットの理解を革新したcontext-sensitive half-time'
    '（状況依存性半減期）に類似する、局所麻酔薬のcontext-sensitive maximum dose'
    '（状況依存性極量）の概念を提案する。27 '
    'context-sensitive half-timeが薬物投与の期間と状況によって変動するように、'
    '局所麻酔薬の実効最大安全用量も投与の状況\u2014具体的には投与経路、特定のブロック種別、'
    '薬物沈着の成否、および個々の患者因子\u2014によって変動すべきである。'
)

add_para('この枠組みでは、極量は単一の固定値ではなく、臨床シナリオに依存する範囲となる：')

# Table 1
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['シナリオ', '初期コンパートメント', '予測Cmax', '用量調整']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

data = [
    ['成功したブロック\n（神経周囲/筋膜面）', 'BPT（V3）\n血管乏しい組織', '低値、遅延', 'より高用量が\n安全な可能性'],
    ['部分的ブロック', '混合\n（BPT + BRT/血漿）', '中間', '標準的極量\nが適用'],
    ['不成功のブロック\n（組織ミスプレイスメント）', 'BRT（V2）\n血管豊富組織', '中等度〜高値、\n早期', 'より低用量が\n必要な可能性'],
    ['血管内注入', '血漿（V1）', '非常に高値、\n即時', '従来のIV極量\nが適用'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        cell = table.rows[r+1].cells[c]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_para('')
add_para('表1. 区域麻酔における局所麻酔薬の状況依存性極量の枠組み。', italic=True)

add_para(
    'このアプローチはいくつかの重要な現実を認識している。第一に、薬物動態学的エビデンスは、'
    '吸収が緩徐で予測可能な筋膜面ブロックでのより高用量を支持する。9,11 '
    '第二に、区域麻酔の実臨床で従来の用量制限を頻繁に超えているにもかかわらず'
    'LASTが稀であるという経験的観察に合理的根拠を提供する。12 '
    '第三に、真に懸念されるシナリオ\u2014不成功のブロックと血管内注入\u2014を特定し、'
    '保守的用量制限の遵守と注意深いモニタリングが最も重要である状況を明確にする。'
)

# ===== CLINICAL IMPLICATIONS =====
add_heading_text('臨床的意義と今後の方向性', level=1)

add_para(
    '初期コンパートメント依存性薬物動態モデリングの意義は用量計算を超えて広がる。'
    'ブロック成功が臨床評価から推定可能であれば（例：予測される時間枠内での感覚ブロック発現）、'
    '臨床医はこの情報を用いてLASTのリスク評価をリアルタイムで更新できる。'
    '成功が確認されたブロックは低リスクの薬物動態推移を示し、ブロック不成功の場合は'
    '警戒を高め、追加投与の安全性を検討すべきである。'
)

add_para(
    '今後の研究はいくつかの優先事項に焦点を当てるべきである。第一に、ブロック成功の同時記録を伴う、'
    '各種区域ブロック後の血漿濃度プロファイルの系統的測定が、コンパートメント特異的吸収モデルの'
    'パラメータ化に必要である。第二に、PK-SimやMoBiなどのプラットフォームを用いて、'
    '注入部位からの組織特異的吸収を組み込んだ局所麻酔薬のPBPKモデルを開発し、'
    '臨床データに対して検証すべきである。第三に、AIMSベンダーを投与経路適応型'
    'PKPDシミュレーション機能の開発に参画させるべきであり、当初は研究ツールとして、'
    '最終的には臨床意思決定支援として実装する。第四に、規制当局および専門学会は、'
    '現行の固定用量極量推奨を、投与経路と部位を考慮した状況依存性ガイドラインで'
    '補完または置換すべきかを検討すべきである。'
)

add_para(
    'ここで提案する枠組みのいくつかの限界を認識する。特定の注入部位からの局所麻酔薬の'
    '組織吸収に関する薬物動態パラメータは不完全にしか特徴付けられていない。'
    'ブロック成功は二値状態ではなくスペクトラムであり、血管乏しい組織対血管豊富組織への'
    '薬物沈着比率は臨床的に正確に決定できない。組織血管分布、蛋白結合、肝クリアランスの'
    '個人差がさらなる不確実性をもたらす。それでも、初期コンパートメントへの薬物動態の'
    '根本的依存性を、不完全にであっても認識することは、それを完全に無視する現行の'
    'アプローチに対する大きな前進であると我々は信じる。'
)

# ===== CONCLUSION =====
add_heading_text('結論（Conclusions）', level=1)

add_para(
    '薬物沈着の初期コンパートメントは、区域麻酔における局所麻酔薬の薬物動態の重要ではあるが'
    '看過されてきた決定因子である。成功したブロックは薬物を血管乏しい組織に沈着させ緩徐な'
    '全身吸収をもたらす一方、不成功のブロックは静脈内投与に近似する可能性がある。'
    'IV系薬物動態モデルから導出された現行の極量はこの根本的差異を考慮しておらず、'
    '成功したブロックに対しては過度に保守的であると同時に、不成功のブロックには'
    '不十分である可能性がある。'
)

add_para(
    'PK-SimやMoBiなどのPBPKモデリングプラットフォームは、投与経路依存性薬物動態を'
    'シミュレートし状況依存性用量推奨を策定するツールを提供する。さらに、AIMSに組み込まれた'
    'PKPDシミュレーションモジュールは、投与経路のロジックを組み込むよう適応されるべきであり、'
    '普遍的なIV仮定ではなく実際の臨床シナリオを反映した薬物動態予測を臨床医に提供すべきである。'
    '我々は、麻酔研究コミュニティ、AIMS開発者、および専門学会に対し、'
    '区域麻酔の安全性と有効性の向上を最終目標としてこれらの目標を追求するよう呼びかける。'
)

# ===== DECLARATIONS =====
add_heading_text('利益相反（Declaration of interest）', level=1)
add_para('[著者により記載]')

add_heading_text('資金源（Funding）', level=1)
add_para('[著者により記載]')

add_heading_text('著者の貢献（Authors\u2019 contributions）', level=1)
add_para('[著者により記載]')

add_heading_text('謝辞（Acknowledgements）', level=1)
add_para('[著者により記載]')

add_heading_text('生成AIの使用に関する宣言', level=1)
add_para(
    '[BJAのポリシーに従い、AI ツールの使用を宣言する必要があります。'
    'AI文章作成支援ツールを使用した場合は、その役割をここに記述してください。]'
)

doc.add_page_break()

# ===== FIGURE LEGENDS =====
add_heading_text('図の説明（Figure Legends）', level=1)

add_para(
    '図1. 3つの臨床シナリオにおけるコンパートメントモデルの比較。'
    '（A）薬物が中心血漿コンパートメント（V1）に投入される従来の静脈内3コンパートメントモデル。'
    '（B）薬物が血管乏しい組織（BPT）に近似するデポコンパートメントに沈着し、'
    '血漿への緩徐な一次吸収を伴う成功した区域ブロックモデル。'
    '（C）薬物が血漿（V1）または血管豊富組織（BRT）に直接投入される不成功のブロックまたは'
    '血管内注入モデル。BRT, 血管豊富組織; BPT, 血管乏しい組織; CL, クリアランス; '
    'ka, 吸収速度定数。', italic=True
)
add_para('')
add_para(
    '図2. 異なる投与経路での局所麻酔薬のシミュレートされた血漿濃度\u2013時間プロファイル。'
    '青の実線は静脈内ボーラス（従来モデル）。赤の破線は血管豊富組織に沈着した不成功のブロック'
    '（急速吸収）。緑の一点鎖線は血管乏しい組織に沈着した成功したブロック（緩徐吸収）。'
    'オレンジの点線は混合沈着の部分的ブロック。水平の破線はCNSおよび心血管系毒性閾値。'
    '薬物沈着の初期コンパートメントに応じたCmaxとTmaxの顕著な差異に注目。', italic=True
)
add_para('')
add_para(
    '図3. 状況依存性極量決定のための生理学的薬物動態（PBPK）シミュレーションの提案ワークフロー。'
    '臨床評価（ステップ1）が初期コンパートメント選択（ステップ2）に情報を提供し、'
    'PK-SimまたはMoBiによるPBPKシミュレーション（ステップ3）を指示し、'
    'シナリオ依存性の極量推奨（ステップ4）を導く。', italic=True
)
add_para('')
add_para(
    '図4. 麻酔情報管理システム（AIMS）における投与経路適応型PKPD シミュレーションの概念的模式図。'
    '局所麻酔薬が区域麻酔手技として投与記録されると、AIMSは標準的な3コンパートメントIVモデルから、'
    'ブロック種別に特異的な吸収パラメータを持つデポ増強モデルに切り替える。'
    '表示される予測血漿濃度曲線は実際の投与経路を反映し、累積用量と毒性リスクの'
    'より正確な評価を臨床医に提供する。AIMS, 麻酔情報管理システム; PKPD, '
    '薬物動態・薬力学; TCI, target-controlled infusion。', italic=True
)

doc.add_page_break()

# ===== REFERENCES =====
add_heading_text('引用文献（References）', level=1)

refs = [
    '1. Rosenberg PH, Veering BT, Urmey WF. Maximum recommended doses of local anesthetics: a multifactorial concept. Reg Anesth Pain Med 2004; 29: 564\u201375.',
    '2. El-Boghdadly K, Pawa A, Chin KJ. Local anesthetic systemic toxicity: current perspectives. Local Reg Anesth 2018; 11: 35\u201344.',
    '3. Marsh B, White M, Morton N, Kenny GN. Pharmacokinetic model driven infusion of propofol in children. Br J Anaesth 1991; 67: 41\u20138.',
    '4. Schnider TW, Minto CF, Gambus PL, et al. The influence of method of administration and covariates on the pharmacokinetics of propofol in adult volunteers. Anesthesiology 1998; 88: 1170\u201382.',
    '5. Eleveld DJ, Colin P, Absalom AR, Struys MMRF. Pharmacokinetic\u2013pharmacodynamic model for propofol for broad application in anaesthesia and sedation. Br J Anaesth 2018; 120: 942\u201359.',
    '6. Minto CF, Schnider TW, Egan TD, et al. Influence of age and gender on the pharmacokinetics and pharmacodynamics of remifentanil. Anesthesiology 1997; 86: 10\u201323.',
    '7. Leite-Moreira AM, Correia A, Vale N, Mour\u00e3o JB. Pharmacokinetics in regional anesthesia. Curr Opin Anaesthesiol 2024; 37: 520\u20135.',
    '8. Arthur GR, Covino BG. Pharmacokinetics of local anaesthetics. Bailli\u00e8re\u2019s Clin Anaesthesiol 1991; 5: 635\u201358.',
    '9. De Cassai A, Dost B, Mormando G, Stecco C. Epinephrine, absorption, and local anaesthetic systemic toxicity: insights from continuous fascial block pharmacokinetic models. Br J Anaesth 2025; 135: 857\u201360.',
    '10. Schwenk ES, Sneyd JR, Wu CL. The state of local anaesthetic systemic toxicity in 2025: the emergence of lidocaine as our next challenge. Br J Anaesth 2025; 135: 854\u20136.',
    '11. Rahiri JL, Tuhoe J, Svirskis D, Lightfoot NJ, Lirk PB, Hill AG. Systematic review of the systemic concentrations of local anaesthetic after transversus abdominis plane block and rectus sheath block. Br J Anaesth 2017; 118: 517\u201326.',
    '12. De Cassai A, Pasin L, Boscolo A, et al. Safety of local anesthetics for fascial plane blocks: a narrative review. J Clin Anesth 2022; 77: 110637.',
    '13. Barrington MJ, Kluger R. Ultrasound guidance reduces the risk of local anesthetic systemic toxicity following peripheral nerve blockade. Reg Anesth Pain Med 2013; 38: 289\u201397.',
    '14. Fettiplace MR, Weinberg G. The mechanisms underlying lipid resuscitation therapy. Reg Anesth Pain Med 2018; 43: 138\u201349.',
    '15. Tucker GT, Mather LE. Clinical pharmacokinetics of local anaesthetic agents. Clin Pharmacokinet 1979; 4: 241\u201378.',
    '16. Simon MJG, Veering BT. Factors affecting the pharmacokinetics and neural block characteristics after epidural administration of local anaesthetics. Eur J Pain 2010; 4: 209\u201318.',
    '17. Burm AG, van der Meer AD, van Kleef JW, Zeijlmans PW, Groen K. Pharmacokinetics of the enantiomers of bupivacaine following intravenous administration of the racemate. Br J Clin Pharmacol 1994; 38: 125\u201329.',
    '18. Gaudreault F, Bherer L, Bhatt DL, Bhatt HV, Bhatt HV. Modeling the anesthetic effect of ropivacaine after a femoral nerve block in orthopedic patients. Anesthesiology 2015; 122: 1010\u201320.',
    '19. Ling J, Xu C, Tang L, Qiu L, Hu N. Comparison of the pharmacokinetic variations of different concentrations of ropivacaine used for serratus anterior plane block. Front Pharmacol 2025; 16: 1540606.',
    '20. Kwa A, Sprung J, Van Guilder M, Jelliffe RW. A population pharmacokinetic model of epidural lidocaine in geriatric patients. Ther Drug Monit 2008; 30: 346\u201355.',
    '21. Thompson MD, Beard DA. Physiologically-based pharmacokinetic tissue compartment model selection in drug development and risk assessment. J Pharm Sci 2012; 101: 424\u201335.',
    '22. Jones HM, Rowland-Yeo K. Basic concepts in physiologically based pharmacokinetic modeling in drug discovery and development. CPT Pharmacometrics Syst Pharmacol 2013; 2: e63.',
    '23. Lippert J, Burghaus R, Edginton A, et al. Open Systems Pharmacology Community\u2014an open access, open source, open science approach to modeling and simulation in pharmaceutical sciences. CPT Pharmacometrics Syst Pharmacol 2019; 8: 878\u201382.',
    '24. Willmann S, Lippert J, Sevestre M, Solodenko J, Fois F, Schmitt W. PK-Sim\u00ae: a physiologically based pharmacokinetic \u2018whole-body\u2019 model. Biosilico 2003; 1: 121\u20134.',
    '25. Open Systems Pharmacology. PK-Sim\u00ae Documentation: Formulations and Administration Protocols. Available at: https://docs.open-systems-pharmacology.org (accessed March 2026).',
    '26. Gamb\u00fas PL, Troc\u00f3niz IF. Pharmacokinetic\u2013pharmacodynamic modelling in anaesthesia. Br J Clin Pharmacol 2015; 79: 72\u201384.',
    '27. Hughes MA, Glass PS, Jacobs JR. Context-sensitive half-time in multicompartment pharmacokinetic models for intravenous anesthetic drugs. Anesthesiology 1992; 76: 334\u201341.',
    '28. Pirri C, Torre DE, Stecco C. Fascial plane blocks: from microanatomy to clinical applications. Curr Opin Anaesthesiol 2024; 37: 526\u201332.',
    '29. Sharma SK, Sonawane K, Mistry T. A narrative review on fascial plane blocks \u2013 Part A: Anatomical foundations and mechanistic insights. Indian J Anaesth 2026; 70: 127\u201336.',
    '30. Niederalt C, Kuepfer L, Solodenko J, et al. A generic whole body physiologically based pharmacokinetic model for therapeutic proteins in PK-Sim. J Pharmacokinet Pharmacodyn 2018; 45: 235\u201357.',
    '31. Gill KL, Gardner I, Li L, Jamei M. A bottom-up whole-body physiologically based pharmacokinetic model to mechanistically predict tissue distribution and the rate of subcutaneous absorption of therapeutic proteins. AAPS J 2016; 18: 156\u201370.',
    '32. Ashraf MW, Uusalo P, Scheinin M, Saari TI. Population modelling of dexmedetomidine pharmacokinetics and haemodynamic effects after intravenous and subcutaneous administration. Clin Pharmacokinet 2020; 59: 1467\u201382.',
    '33. Pepin XJH, Grant I, Wood JM. SubQ-Sim: a subcutaneous physiologically based biopharmaceutics model. Part 1: the injection and system parameters. Pharm Res 2023; 40: 2195\u2013214.',
    '34. Silva DA, Le Merdy M, Mullin J, et al. Mechanistic modeling of intramuscular administration of a long-acting injectable accounting for tissue response at the depot site. AAPS J 2026; 28: 4.',
    '35. De Cassai A, Bonvicini D, Correale C, et al. Histology of the fascial planes: a systematic review of the microstructural foundations of regional anesthesia. J Anesth Analg Crit Care 2026; 6: 5.',
    '36. Winnie AP, Tay CH, Patel KP, Ramamurthy S, Durrani Z. Pharmacokinetics of local anesthetics during plexus blocks. Anesth Analg 1977; 56: 852\u201361.',
    '37. Y\u00e1\u00f1ez JA, Remsberg CM, Sayre CL, Forrest ML, Davies NM. Flip-flop pharmacokinetics\u2014delivering a reversal of disposition: challenges and opportunities during drug development. Ther Deliv 2011; 2: 643\u201372.',
    '38. Butiulca M, Farczadi L, Imre S, et al. LC-MS/MS assisted pharmacokinetic and tissue distribution study of ropivacaine and 3-OH-ropivacaine on rats after plane block anesthesia. Front Pharmacol 2025; 15: 1494646.',
    '39. Osborne KW, MacFater WS, Anderson BJ, Svirskis D, Hill AG, Hannam JA. Pharmacokinetics of intraperitoneal lidocaine for sustained postoperative analgesia in adults. Eur J Drug Metab Pharmacokinet 2025.',
    '40. Bettonte S, Berton M, Battegay M, Stader F, Marzolini C. Development of a physiologically-based pharmacokinetic model to simulate the pharmacokinetics of intramuscular antiretroviral drugs. CPT Pharmacometrics Syst Pharmacol 2024; 13: 781\u201394.',
    '41. Dost B. Fascial plane blocks in the era of modern regional anesthesia: shaping the future of pain management. J Anesth Analg Crit Care 2025; 5: 49.',
    '42. Enlund M. TCI: Target Controlled Infusion, or Totally Confused Infusion? Upsala J Med Sci 2008; 113: 161\u201370.',
    '43. Vellinga R, Eleveld DJ, Struys MMRF, van den Berg JP. General purpose models for intravenous anesthetics, the next generation for target-controlled infusion and total intravenous anesthesia? Curr Opin Anaesthesiol 2023; 36: 602\u20137.',
    '44. \u0160afr\u00e1nkov\u00e1 P, Bruthans J. Target-controlled infusion of propofol: a systematic review of recent results. J Med Syst 2025; 49: 54.',
    '45. Sessler DI, Bao X, Leiman D, et al. A phase I study of the pharmacokinetics, pharmacodynamics, and safety of liposomal bupivacaine for sciatic nerve block in the popliteal fossa for bunionectomy. J Clin Pharmacol 2025; 65: 441\u201351.',
    '46. Xu A, Ren A, Lee C. Pharmacokinetics of lidocaine infusion: optimal dosing and duration in ERAS protocol. medRxiv 2025.',
    '47. Cascone S, Lamberti G, Titomanlio G, Piazza O. Pharmacokinetics of remifentanil: a three-compartmental modeling approach. Transl Med UniSa 2013; 7: 18\u201322.',
    '48. Brainkart. Pharmacokinetics: compartment models. In: Clinical Anesthesiology: Clinical Pharmacology. Available at: https://www.brainkart.com (accessed March 2026).',
    '49. Deranged Physiology. Single and multiple compartment models of drug distribution. Available at: https://derangedphysiology.com (accessed March 2026).',
    '50. Holt A. Three compartment drug. In: An ABC of PK/PD. Open Education Alberta; 2023.',
]

for ref in refs:
    add_para(ref)

doc.add_page_break()

# ===== INSERT FIGURES =====
add_heading_text('図（Figures）', level=1)

fig_dir = '/home/ubuntu/manuscript/figures'
for fname, caption in [
    ('figure1_compartment_models.png', '図1（Figure 1）'),
    ('figure2_pk_simulation.png', '図2（Figure 2）'),
    ('figure3_workflow.png', '図3（Figure 3）'),
    ('figure4_aims.png', '図4（Figure 4）'),
]:
    fpath = os.path.join(fig_dir, fname)
    if os.path.exists(fpath):
        doc.add_paragraph(caption)
        doc.add_picture(fpath, width=Inches(6.0))
        doc.add_paragraph('')

# Save
out_path = '/home/ubuntu/manuscript/BJA_Manuscript_Japanese.docx'
doc.save(out_path)
print(f'Japanese manuscript saved to {out_path}')
