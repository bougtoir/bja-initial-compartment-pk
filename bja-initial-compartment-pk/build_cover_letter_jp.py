#!/usr/bin/env python3
"""Generate Japanese BJA cover letter as .docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

def add_para(text, bold=False, italic=False, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ===== DATE =====
add_para('[日付]', space_after=12)

# ===== ADDRESSEE =====
add_para('Professor Hugh Hemmings')
add_para('Editor-in-Chief')
add_para('British Journal of Anaesthesia')
add_para('Weill Cornell Medical College')
add_para('New York, NY, USA', space_after=12)

# ===== SUBJECT =====
add_para('Dear Professor Hemmings,', space_after=12)

add_para(
    'Re: Narrative Review投稿 \u2014 '
    '\u201cRethinking Maximum Dose Limits for Local Anaesthetics in Regional Anaesthesia: '
    'The Case for Initial Compartment-Dependent Pharmacokinetic Modelling and a Call for '
    'Route-Adaptive PKPD Simulation in Anaesthesia Information Management Systems\u201d',
    bold=True, space_after=12
)

add_para(
    '（区域麻酔における局所麻酔薬の極量再考：初期コンパートメント依存性薬物動態モデリングの'
    '必要性および麻酔情報管理システムにおける投与経路適応型PKPDシミュレーションの提言）',
    italic=True, space_after=12
)

# ===== BODY =====
add_para(
    '上記原稿をBritish Journal of AnaesthesiaにNarrative Reviewとしてご査読いただきたく、'
    '投稿申し上げます。本稿は、区域麻酔の薬物動態における基本的かつ見過ごされてきた問題'
    '\u2014すなわち、現行の局所麻酔薬極量が静脈内投与を前提としたモデルから導出されている'
    'にもかかわらず、区域麻酔における初期薬物動態コンパートメントは投与経路およびブロックの'
    '成否に決定的に依存する\u2014に取り組むものです。',
    space_after=12
)

add_para(
    '本レビューでは、薬物沈着の初期コンパートメント\u2014全身吸収が緩徐なブロックでは'
    '血管乏しい組織、全身吸収が急速なブロックでは血管豊富組織または血漿\u2014が薬物動態推移、'
    'ひいては局所麻酔薬全身毒性（LAST）のリスクを決定するという概念的枠組みを提案します。'
    'この枠組みは、無料で利用可能な生理学的薬物動態（PBPK）モデリングプラットフォーム'
    'であるPK-SimおよびMoBiを用いて実装可能であることを示すとともに、現代の麻酔情報管理'
    'システム（AIMS）に組み込まれた薬物動態\u2013薬力学（PKPD）シミュレーションモジュールが'
    'リアルタイム用量ガイダンスのために投与経路ロジックを組み込むべきであると提言します。',
    space_after=12
)

add_para(
    '本稿は、De Cassaiらによる最近のBJA論説（2025年）\u2014筋膜面ブロック後の局所麻酔薬の'
    '母集団薬物動態研究の必要性を提唱した\u2014で紹介された概念を発展・拡張するものです。'
    '本研究が、麻酔科医、薬理学者、臨床意思決定支援システム開発者を含むBJA読者にとって'
    '時宜を得た包括的な枠組みを提供するものと確信しております。',
    space_after=12
)

add_para(
    '原稿は本文約4800語、引用文献50本、図4枚（シミュレーション由来の図2枚と概念的フロー図2枚）、'
    '表1枚で構成されています。これらはBJA Narrative Reviewの制限（5000語、150文献、'
    '図表合計6点）の範囲内です。全ての図はカラーで300 dpi以上の解像度で提供しております。',
    space_after=12
)

add_para(
    '本原稿はオリジナルの著作であり、他誌に掲載されたことはなく、現在他誌で査読中でも'
    'ありません。全著者が最終原稿を読了・承認し、British Journal of Anaesthesiaへの'
    '投稿に同意しております。',
    space_after=12
)

add_para(
    '本研究に関連する利益相反はありません。'
    '［該当する場合は、関連する利益相反を開示するようこの記述を修正してください。］',
    space_after=12
)

add_para(
    '査読者のフィードバックに基づき原稿を修正する機会をいただければ幸いです。'
    'ご検討のほどよろしくお願い申し上げます。',
    space_after=12
)

add_para('敬具', space_after=24)

add_para('[責任著者名]')
add_para('[役職、所属部門]')
add_para('[所属機関]')
add_para('[住所]')
add_para('[メールアドレス]')
add_para('[電話番号]')

# ===== SAVE =====
out = os.path.join(os.path.dirname(__file__), 'BJA_Cover_Letter_Japanese.docx')
doc.save(out)
print(f'Japanese cover letter saved to {out}')
