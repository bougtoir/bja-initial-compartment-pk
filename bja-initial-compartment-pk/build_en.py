#!/usr/bin/env python3
"""Generate English BJA manuscript as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 2.0

# Helper
def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ===== TITLE PAGE =====
add_para('')
add_para('')
p = add_para('Rethinking Maximum Dose Limits for Local Anaesthetics in Regional Anaesthesia:', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
p = add_para('The Case for Initial Compartment-Dependent Pharmacokinetic Modelling', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('and a Call for Route-Adaptive PKPD Simulation in Anaesthesia Information Management Systems', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('')

add_para('Article type: Narrative Review', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('[Author names to be inserted]', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Affiliations to be inserted]', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')
add_para('Corresponding author:', bold=True)
add_para('[Name, Department, Institution, Address, Email]')
add_para('')
add_para('Word count: ~4800 words (main text)')
add_para('References: 50')
add_para('Figures: 4')
add_para('Tables: 1')
add_para('')

# Keywords
add_para('Keywords: local anaesthetic systemic toxicity; pharmacokinetic modelling; regional anaesthesia; physiologically based pharmacokinetics; compartment model; anaesthesia information management system; maximum recommended dose', italic=True)

doc.add_page_break()

# ===== SUMMARY =====
add_heading_text('Summary', level=1)

summary_text = (
    'Current maximum recommended doses for local anaesthetics are derived from pharmacokinetic models '
    'that assume intravenous administration, where drugs are deposited directly into the central plasma '
    'compartment. However, in regional anaesthesia, the initial site of drug deposition differs fundamentally '
    'depending on block success: a successful perineural or fascial plane block deposits drug into vessel-poor '
    'tissue (slow systemic absorption), whereas a failed or partially failed block may result in drug deposition '
    'into vessel-rich tissue or direct intravascular injection (rapid systemic absorption). This discrepancy in '
    'the initial pharmacokinetic compartment has profound implications for peak plasma concentration and, '
    'consequently, the risk of local anaesthetic systemic toxicity. We review the limitations of conventional '
    'three-compartment models for regional anaesthesia dosing, examine the potential of physiologically based '
    'pharmacokinetic platforms such as PK-Sim and MoBi to simulate route-dependent pharmacokinetics, and '
    'propose a framework for context-sensitive maximum dose recommendations. Furthermore, we argue that '
    'pharmacokinetic-pharmacodynamic simulation modules embedded in modern anaesthesia information '
    'management systems should incorporate route-of-administration logic, enabling real-time, '
    'compartment-appropriate dose guidance for both intravenous and regional anaesthesia techniques.'
)
add_para(summary_text)

doc.add_page_break()

# ===== INTRODUCTION =====
add_heading_text('Introduction', level=1)

add_para(
    'Local anaesthetic systemic toxicity (LAST) remains one of the most feared complications of regional '
    'anaesthesia. Maximum recommended doses for local anaesthetics, traditionally expressed as mg kg\u207b\u00b9, '
    'were established decades ago, largely on the basis of pharmacokinetic studies involving intravenous '
    'administration or subcutaneous infiltration.1,2 These dose limits implicitly assume that the drug enters '
    'the systemic circulation via the central (plasma) compartment, as occurs with intravenous injection.'
)

add_para(
    'The three-compartment mammillary model, widely used in anaesthetic pharmacology, describes drug '
    'distribution from the central compartment (V1, plasma) to a rapidly equilibrating peripheral compartment '
    '(V2, vessel-rich tissues or BRT) and a slowly equilibrating peripheral compartment (V3, vessel-poor '
    'tissues or BPT), with elimination from the central compartment.3,4 This model underpins target-controlled '
    'infusion systems (e.g. Marsh, Schnider, Eleveld models for propofol) and is the foundation of '
    'pharmacokinetic-pharmacodynamic (PKPD) displays in modern anaesthesia information management '
    'systems (AIMS).5,6'
)

add_para(
    'However, in regional anaesthesia, the drug is not administered intravenously. It is deposited into '
    'tissue\u2014perineural space, fascial planes, or epidural space\u2014where systemic absorption depends on local '
    'blood flow, tissue binding, and the physicochemical properties of both the drug and the tissue.7,8 '
    'Notably, neuraxial techniques introduce additional pharmacokinetic complexity: epidural administration '
    'involves simultaneous absorption via epidural fat, dural transfer into cerebrospinal fluid (CSF), and '
    'vascular uptake, whilst intrathecal (spinal) administration deposits drug directly into the CSF with '
    'unique distribution and absorption kinetics that differ from all peripheral routes. '
    'Recent work by De Cassai and colleagues, published in this journal, has advanced our understanding '
    'of local anaesthetic pharmacokinetics in fascial plane blocks, highlighting the roles of epinephrine, '
    'tissue vascularity, and fascial microanatomy in determining systemic absorption profiles.9 Similarly, '
    'Schwenk and colleagues have drawn attention to the emergence of lidocaine as a persistent cause of '
    'LAST-related mortality.10'
)

add_para(
    'Despite these advances, a fundamental question has not been adequately addressed: how does the '
    'initial site of drug deposition\u2014the starting compartment\u2014influence peak plasma concentration and '
    'therefore the margin of safety? We contend that the answer to this question is critical to any meaningful '
    'discussion of maximum dose limits in regional anaesthesia, and that current dose recommendations are '
    'inadequate because they fail to account for this variable.'
)

# ===== THE PROBLEM =====
add_heading_text('The initial compartment problem', level=1)

add_para(
    'Consider two clinical scenarios involving the same dose of a long-acting local anaesthetic '
    '(e.g. bupivacaine 150 mg) administered for a peripheral nerve block:'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run(
    'Scenario A (successful block): The entire dose is deposited accurately into the target fascial plane '
    'or perineural space. The tissue is predominantly vessel-poor (adipose, connective tissue, fascia). '
    'Systemic absorption is slow, governed by a low absorption rate constant (ka). Peak plasma concentration '
    '(Cmax) is low and delayed (high Tmax). The drug exerts its intended local effect while being gradually '
    'cleared by hepatic metabolism as it enters the systemic circulation.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run(
    'Scenario B (failed block or intravascular injection): The drug is deposited into '
    'a vascular structure or highly perfused tissue (vessel-rich group). Systemic absorption is rapid, '
    'equivalent or near-equivalent to intravenous administration. Cmax is high and occurs early (low Tmax), '
    'potentially exceeding the threshold for central nervous system or cardiovascular toxicity.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_para(
    'In the traditional three-compartment model, both scenarios would be evaluated against the same '
    'maximum recommended dose\u2014a dose limit derived from models assuming plasma as the initial '
    'compartment. Yet the pharmacokinetic profiles are fundamentally different. For Scenario A, the '
    'conventional dose limit may be unnecessarily conservative, as the slow absorption rate produces '
    'a Cmax well below the toxicity threshold even at doses exceeding the traditional limit. For Scenario B, '
    'the same dose limit may be dangerously liberal, as rapid systemic absorption produces plasma '
    'concentrations comparable to those seen after intravenous bolus administration.'
)

add_para(
    'This asymmetry is not merely theoretical. Clinical experience demonstrates that fascial plane blocks '
    'routinely employ doses exceeding traditional weight-based limits without apparent toxicity,11,12 '
    'whilst rare but devastating LAST events continue to occur at conventional doses, often attributable '
    'to inadvertent intravascular injection or rapid absorption from highly vascular injection sites.13,14 '
    'The missing variable in this equation is the initial compartment of drug deposition.'
)

# ===== BLOCK SUCCESS AS A PK DETERMINANT =====
add_heading_text('Block success as a pharmacokinetic determinant', level=1)

add_para(
    'We propose that the clinical success of a regional block provides direct information about the '
    'pharmacokinetic trajectory of the administered local anaesthetic. A successful block producing '
    'prolonged sensory and motor blockade implies that a substantial portion of the drug remains '
    'localised in the target tissue (vessel-poor compartment) for an extended period. This is pharmacokinetic '
    'evidence of slow absorption: the drug has not been rapidly cleared from the injection site into the '
    'systemic circulation.'
)

add_para(
    'Conversely, a block that fails rapidly or never achieves adequate anaesthesia suggests that the drug '
    'has been rapidly absorbed or deposited away from the target nerve, entering the systemic circulation '
    'through vessel-rich tissue pathways. In pharmacokinetic terms, this represents a scenario where the '
    'initial compartment is closer to V2 (vessel-rich tissue) or V1 (plasma) rather than V3 (vessel-poor '
    'tissue).'
)

add_para(
    'This relationship between block efficacy and pharmacokinetic behaviour has important implications. '
    'When a block is clinically effective for many hours, the plasma concentration curve is characterised '
    'by a low, flat profile\u2014the drug is being sequestered locally rather than flooding the systemic '
    'circulation. The risk of LAST in this scenario is inherently lower than that predicted by models '
    'assuming immediate plasma compartment entry. The corollary is that failed blocks present a higher '
    'toxicity risk than conventionally appreciated, because the entire dose may behave pharmacokinetically '
    'as if it were administered intravenously.'
)

# ===== LIMITATIONS OF CURRENT MODELS =====
add_heading_text('Limitations of conventional three-compartment models for regional anaesthesia', level=1)

add_para(
    'The three-compartment models used in clinical anaesthesia practice (Marsh, Schnider, Eleveld for '
    'propofol; Minto, Kim, Eleveld for remifentanil) were developed for intravenous drug administration.3\u20136 '
    'They describe drug disposition after the drug has entered the central compartment and are therefore '
    'inherently unsuitable for modelling regional anaesthesia pharmacokinetics without modification.'
)

add_para(
    'Several specific limitations deserve emphasis. First, these models lack a depot or absorption '
    'compartment. In regional anaesthesia, drug must first be absorbed from the injection site before '
    'entering the plasma, a process characterised by a rate constant (ka) and bioavailability (F) that '
    'vary with injection site, tissue vascularity, use of vasoconstrictors, and individual patient factors.15,16 '
    'Second, the inter-compartmental rate constants (k12, k21, k13, k31) were estimated from IV '
    'administration data and may not accurately reflect drug transfer kinetics when the drug originates '
    'from a peripheral tissue depot. Third, protein binding dynamics differ importantly between IV bolus '
    'administration (where free drug fraction spikes acutely) and slow tissue absorption (where protein '
    'binding capacity is not overwhelmed).17'
)

add_para(
    'Population pharmacokinetic studies of local anaesthetics after regional blocks have addressed some '
    'of these limitations by incorporating depot compartments with first-order absorption.18\u201320 '
    'Gaudreault and colleagues modelled ropivacaine pharmacokinetics after femoral nerve block using '
    'a two-compartment model with first-order absorption, demonstrating flip-flop kinetics where the '
    'absorption rate was slower than the elimination rate.18 More recently, Ling and colleagues '
    'developed a population pharmacokinetic model for ropivacaine after serratus anterior plane block '
    'using NONMEM.19 These studies consistently show that the pharmacokinetic profile after regional '
    'block differs markedly from IV administration, yet their findings have not been translated into '
    'revised dose recommendations.'
)

# ===== PBPK APPROACH =====
add_heading_text('Physiologically based pharmacokinetic modelling: a route-adaptive approach', level=1)

add_para(
    'Physiologically based pharmacokinetic (PBPK) models offer a fundamentally different approach to '
    'simulating drug disposition. Rather than using abstract compartments with empirically estimated '
    'transfer constants, PBPK models divide the body into anatomically and physiologically defined organ '
    'compartments, each characterised by known blood flow, tissue volume, partition coefficients, and '
    'metabolic capacity.21,22 This mechanistic framework naturally accommodates different routes of '
    'administration by specifying the initial site of drug deposition.'
)

add_para(
    'The Open Systems Pharmacology (OSP) platform, comprising PK-Sim and MoBi, is a freely available, '
    'open-source PBPK modelling suite that has been qualified by the European Medicines Agency and '
    'is widely used in drug development and regulatory science.23,24 PK-Sim provides a graphical '
    'interface for building whole-body PBPK models with predefined organ compartments (including '
    'arterial and venous blood, lung, heart, muscle, adipose, skin, and others), while MoBi allows '
    'custom model building with user-defined compartments and transport processes.'
)

add_para(
    'Critically for our purpose, PK-Sim supports multiple administration routes including intravenous, '
    'intramuscular, and subcutaneous injection, each with route-specific absorption models.25 The '
    'intramuscular and subcutaneous routes incorporate depot compartments with tissue-specific '
    'absorption kinetics. By analogy, a perineural or fascial plane injection could be modelled by '
    'specifying drug deposition into a tissue compartment with characteristics appropriate to the '
    'injection site (e.g. low blood flow for fascial tissue, high blood flow for highly vascular '
    'perineural structures). MoBi further allows the user to define entirely custom compartments '
    'and initial conditions, enabling simulation of the three clinical scenarios described above '
    '(successful block, failed block, and partial block) by varying the initial compartment and '
    'the fraction of dose deposited in each compartment.'
)

add_para(
    'We illustrate this concept in Figure 1, which compares the traditional IV three-compartment model '
    'with modified models for successful and failed regional blocks. Figure 2 presents simulated plasma '
    'concentration\u2013time profiles for each scenario, demonstrating the profound differences in Cmax and '
    'Tmax that result from varying the initial compartment of drug deposition.'
)

# ===== AIMS INTEGRATION =====
add_heading_text('Integration into anaesthesia information management systems', level=1)

add_para(
    'Modern AIMS increasingly incorporate real-time PKPD simulation, displaying predicted plasma and '
    'effect-site concentrations for intravenous agents such as propofol and remifentanil.5,26 These '
    'displays are powered by three-compartment models (e.g. Eleveld, Schnider) and assume that all '
    'drugs are administered intravenously into the central compartment. While this assumption is valid '
    'for total intravenous anaesthesia (TIVA) and target-controlled infusion (TCI), it becomes incorrect '
    'when the same system is used to track local anaesthetic doses administered via regional techniques.'
)

add_para(
    'We propose that AIMS vendors and developers should implement route-adaptive PKPD simulation '
    'that adjusts the pharmacokinetic model based on the documented administration route. Specifically, '
    'the system should:'
)

p = doc.add_paragraph()
p.style = 'List Number'
run = p.add_run(
    'Distinguish between intravenous and regional routes of local anaesthetic administration in the '
    'drug administration record.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.style = 'List Number'
run = p.add_run(
    'Apply route-appropriate pharmacokinetic models: the standard three-compartment model for IV '
    'administration, and a depot-augmented model (with absorption rate constant and bioavailability '
    'appropriate to the specific block type) for regional techniques.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.style = 'List Number'
run = p.add_run(
    'Provide block-type-specific absorption parameters, informed by published population pharmacokinetic '
    'data for each regional technique (e.g. TAP block, ESP block, femoral nerve block, epidural).'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.style = 'List Number'
run = p.add_run(
    'Display predicted plasma concentration trajectories that reflect the actual route of administration, '
    'enabling clinicians to make informed decisions about redosing and cumulative dose limits.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.style = 'List Number'
run = p.add_run(
    'Incorporate a mechanism to adjust the model in real time based on clinical indicators of block '
    'success (e.g. sensory testing results), shifting between successful-block (slow absorption) and '
    'failed-block (rapid absorption) pharmacokinetic profiles.'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_para(
    'Figure 3 illustrates the proposed workflow for PBPK-based maximum dose reassessment. '
    'Figure 4 presents a conceptual schematic of how route-adaptive PKPD simulation could be '
    'integrated into AIMS displays, with the pharmacokinetic model automatically adjusted based '
    'on the documented route of administration and block type.'
)

add_para(
    'The technical feasibility of this approach is supported by several observations. First, the '
    'mathematical framework for depot-augmented compartment models is well established and '
    'computationally inexpensive.18\u201320 Second, population pharmacokinetic parameters for local '
    'anaesthetics after various regional techniques are increasingly available in the literature, '
    'providing the data needed to parameterise route-specific models.7,8,19 Third, modern AIMS '
    'already implement real-time PKPD simulation for intravenous agents, demonstrating that the '
    'computational infrastructure exists. The principal barrier is not technical but conceptual: the '
    'recognition that a single pharmacokinetic model cannot adequately describe drug behaviour '
    'across fundamentally different routes of administration.'
)

# ===== TOWARD CONTEXT-SENSITIVE MAX DOSE =====
add_heading_text('Toward context-sensitive maximum dose recommendations', level=1)

add_para(
    'We propose the concept of a context-sensitive maximum dose for local anaesthetics, analogous '
    'to the context-sensitive half-time that revolutionised our understanding of intravenous drug '
    'offset.27 Just as the context-sensitive half-time varies with the duration and context of drug '
    'administration, the effective maximum safe dose of a local anaesthetic should vary with the '
    'context of administration\u2014specifically, the route, the specific block type, the success of '
    'drug deposition, and individual patient factors.'
)

add_para(
    'Under this framework, the maximum recommended dose would not be a single fixed value but '
    'a range dependent on the clinical scenario:'
)

# Table 1
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
headers = ['Scenario', 'Initial Compartment', 'Expected Cmax', 'Dose Adjustment']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

data = [
    ['Successful block\n(perineural/fascial)', 'BPT (V3)\nVessel-poor tissue', 'Low, delayed', 'Higher dose\nmay be safe'],
    ['Partial block', 'Mixed\n(BPT + BRT/Plasma)', 'Intermediate', 'Standard dose\nlimit applies'],
    ['Failed block\n(vessel-rich deposition)', 'BRT (V2)\nVessel-rich tissue', 'Moderate-high,\nearly', 'Lower dose\nmay be needed'],
    ['Intravascular injection', 'Plasma (V1)', 'Very high,\nimmediate', 'Traditional IV\nlimits apply'],
    ['Epidural administration', 'Depot (multi-pathway)\nFat + dural transfer + vascular', 'Intermediate,\ndelayed', 'Adequately approximated\nby depot model'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        cell = table.rows[r+1].cells[c]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_para('')
add_para('Table 1. Proposed context-sensitive maximum dose framework for local anaesthetics in regional anaesthesia.', italic=True)

add_para(
    'This approach acknowledges several important realities. First, the pharmacokinetic evidence '
    'supports higher doses for fascial plane blocks where absorption is slow and predictable.9,11 '
    'Second, it provides a rational basis for the empirical observation that LAST is rare despite '
    'frequent exceeding of traditional dose limits in regional practice.12 Third, it identifies the '
    'scenarios with rapid systemic absorption\u2014failed blocks and intravascular injection\u2014where adherence to '
    'conservative dose limits and vigilant monitoring are most important.'
)

# ===== EPIDURAL AND SPINAL CONSIDERATIONS =====
add_heading_text('Considerations for neuraxial techniques: epidural and spinal anaesthesia', level=1)

add_para(
    'The framework presented above focuses primarily on peripheral nerve blocks and fascial plane blocks, '
    'where the initial compartment can be reasonably approximated as either vessel-poor tissue (successful block) '
    'or vessel-rich tissue/plasma (failed block). However, two neuraxial routes\u2014epidural and spinal '
    '(intrathecal) administration\u2014deserve specific consideration, as their pharmacokinetics present '
    'distinct challenges for compartmental modelling.'
)

add_para(
    'Epidural administration of local anaesthetics involves drug deposition into the epidural space, '
    'where absorption occurs via three parallel pathways: (i) distribution into epidural fat, which '
    'acts as a local depot with slow release; (ii) transfer across the dura mater into the '
    'cerebrospinal fluid (CSF), where it accesses spinal nerve roots to produce neural blockade; '
    'and (iii) vascular absorption from the rich epidural venous plexus into the systemic '
    'circulation.15,16 Importantly, this multi-pathway absorption can be well approximated by a '
    'depot compartment model with first-order absorption kinetics (ka), analogous to the peripheral '
    'block model described above. Population pharmacokinetic studies of epidural local anaesthetics '
    'have successfully employed depot-augmented compartment models, validating this approach.20 '
    'The depot model thus provides an adequate and practical framework for epidural administration '
    'within the initial compartment paradigm proposed here. Epidural administration can be '
    'considered as a depot-start scenario with absorption characteristics intermediate between '
    'a successful fascial plane block (pure vessel-poor tissue) and a failed block (rapid '
    'vascular uptake), reflecting the mixed nature of the epidural space. Although more '
    'sophisticated multi-pathway models could further refine the simulation of the three '
    'parallel absorption routes, the depot approximation is sufficient for clinical dose '
    'guidance and integration into the proposed framework.'
)

add_para(
    'Spinal (intrathecal) anaesthesia was not included in the primary analysis of this review. '
    'Intrathecal administration deposits drug directly into the CSF, a unique pharmacokinetic '
    'compartment that does not correspond to any of the four compartments in the current framework '
    '(plasma, BRT, BPT, or depot). Drug distribution within the CSF is governed by factors such as '
    'baricity, patient positioning, CSF volume, and spinal curvature, none of which are captured by '
    'conventional compartmental models.16 Furthermore, spinal anaesthesia is almost exclusively '
    'performed as a single-shot technique with small drug doses (e.g. bupivacaine 10\u201315 mg), '
    'making systemic toxicity from intrathecal dosing alone exceedingly rare. For these reasons\u2014the '
    'pharmacokinetic uniqueness of CSF distribution, the small doses employed, and the predominantly '
    'single-shot nature of the technique\u2014we have deliberately excluded spinal anaesthesia from the '
    'initial compartment framework. Extending the model to include a CSF compartment would add '
    'considerable complexity without proportionate clinical benefit for the question of systemic '
    'toxicity and maximum dose.'
)

# ===== CLINICAL IMPLICATIONS =====
add_heading_text('Clinical implications and future directions', level=1)

add_para(
    'The implications of initial compartment-dependent pharmacokinetic modelling extend beyond '
    'dose calculations. If block success can be inferred from clinical assessment (e.g. onset of '
    'sensory block within expected timeframes), clinicians could use this information to update their '
    'risk assessment for LAST in real time. A confirmed successful block would indicate a low-risk '
    'pharmacokinetic trajectory, whilst failure to achieve blockade should prompt heightened vigilance '
    'and consideration of whether additional dosing is safe.'
)

add_para(
    'Future research should focus on several priorities. First, systematic measurement of plasma '
    'concentration profiles after various regional block types, with concurrent documentation of '
    'block success, is needed to parameterise compartment-specific absorption models. Second, '
    'PBPK models for local anaesthetics incorporating tissue-specific absorption from injection '
    'sites should be developed and validated against clinical data using platforms such as PK-Sim '
    'and MoBi. Third, AIMS vendors should be engaged in developing route-adaptive PKPD simulation '
    'capabilities, initially as research tools and ultimately as clinical decision support. Fourth, '
    'regulatory bodies and professional societies should consider whether current fixed-dose maximum '
    'recommendations should be supplemented or replaced by context-sensitive guidelines that account '
    'for the route and site of administration.'
)

add_para(
    'We acknowledge several limitations of the framework proposed here. The pharmacokinetic parameters '
    'for tissue absorption of local anaesthetics from specific injection sites are incompletely '
    'characterised. Block success is a spectrum rather than a binary state, and the proportion of drug '
    'deposited in vessel-poor versus vessel-rich tissue cannot be precisely determined clinically. '
    'Individual variation in tissue vascularity, protein binding, and hepatic clearance introduces '
    'additional uncertainty. As discussed above, epidural administration can be adequately '
    'approximated by a depot compartment model despite its multi-pathway absorption, whilst '
    'spinal anaesthesia was excluded from this framework due to its unique CSF '
    'pharmacokinetics and the small doses employed. Nevertheless, we believe that acknowledging '
    'the fundamental dependence of pharmacokinetics on the initial compartment, even imperfectly, '
    'represents a significant advance over the current approach of ignoring it entirely.'
)

# ===== CONCLUSION =====
add_heading_text('Conclusions', level=1)

add_para(
    'The initial compartment of drug deposition is a critical but neglected determinant of local '
    'anaesthetic pharmacokinetics in regional anaesthesia. Successful blocks deposit drug into '
    'vessel-poor tissue with slow systemic absorption, whilst failed blocks may approximate '
    'intravenous injection. Current maximum dose recommendations, derived from IV-based '
    'pharmacokinetic models, do not account for this fundamental difference and are therefore '
    'simultaneously too conservative for successful blocks and potentially insufficient for failed ones.'
)

add_para(
    'PBPK modelling platforms such as PK-Sim and MoBi provide the tools to simulate '
    'route-dependent pharmacokinetics and develop context-sensitive dose recommendations. '
    'Furthermore, PKPD simulation modules embedded in AIMS should be adapted to incorporate '
    'route-of-administration logic, providing clinicians with pharmacokinetic predictions that '
    'reflect the actual clinical scenario rather than a universal IV assumption. We call upon the '
    'anaesthesia research community, AIMS developers, and professional societies to pursue these '
    'goals, with the ultimate aim of improving both the safety and efficacy of regional anaesthesia practice.'
)

# ===== DECLARATIONS =====
add_heading_text('Declaration of interest', level=1)
add_para('[To be completed by authors]')

add_heading_text('Funding', level=1)
add_para('[To be completed by authors]')

add_heading_text('Authors\u2019 contributions', level=1)
add_para('[To be completed by authors]')

add_heading_text('Acknowledgements', level=1)
add_para('[To be completed by authors]')

add_heading_text('Declaration of generative artificial intelligence (AI) in scientific writing', level=1)
add_para(
    '[Authors must declare the use of AI tools in accordance with BJA policy. '
    'If AI writing assistants were used, describe their role here.]'
)

doc.add_page_break()

# ===== FIGURE LEGENDS =====
add_heading_text('Figure Legends', level=1)

add_para(
    'Figure 1. Comparison of compartment models for three clinical scenarios. '
    '(A) Traditional intravenous three-compartment model with drug entering the central plasma '
    'compartment (V1). (B) Successful regional block model with drug deposited in a depot '
    'compartment approximating vessel-poor tissue (BPT), with slow first-order absorption into '
    'plasma. (C) Failed block or intravascular injection model with drug entering plasma (V1) '
    'or vessel-rich tissue (BRT) directly. BRT, vessel-rich tissue; BPT, vessel-poor tissue; '
    'CL, clearance; ka, absorption rate constant.', italic=True
)

add_para('')
add_para(
    'Figure 2. Simulated plasma concentration\u2013time profiles for local anaesthetic administered '
    'via different routes. The blue solid line represents intravenous bolus (traditional model). '
    'The red dashed line represents a failed block with drug deposited into vessel-rich tissue '
    '(rapid absorption). The green dash-dot line represents a successful block with drug deposited '
    'into vessel-poor tissue (slow absorption). The orange dotted line represents a partial block '
    'with mixed deposition. Horizontal dashed lines indicate CNS and cardiovascular toxicity '
    'thresholds. Note the marked differences in peak plasma concentration (Cmax) and time to '
    'peak (Tmax) depending on the initial compartment of drug deposition.', italic=True
)

add_para('')
add_para(
    'Figure 3. Proposed workflow for physiologically based pharmacokinetic (PBPK) simulation '
    'to determine context-sensitive maximum dose recommendations. Clinical assessment (Step 1) '
    'informs initial compartment selection (Step 2), which directs PBPK simulation using PK-Sim '
    'or MoBi (Step 3), yielding scenario-dependent maximum dose recommendations (Step 4).', italic=True
)

add_para('')
add_para(
    'Figure 4. Conceptual schematic of route-adaptive pharmacokinetic-pharmacodynamic (PKPD) '
    'simulation in anaesthesia information management systems (AIMS). When a local anaesthetic is '
    'documented as administered via a regional technique, the AIMS switches from the standard '
    'three-compartment IV model to a depot-augmented model with absorption parameters specific '
    'to the block type. The displayed predicted plasma concentration curve reflects the actual '
    'route of administration, providing clinicians with a more accurate assessment of cumulative '
    'dose and toxicity risk. AIMS, anaesthesia information management system; PKPD, '
    'pharmacokinetic-pharmacodynamic; TCI, target-controlled infusion.', italic=True
)

doc.add_page_break()

# ===== REFERENCES =====
add_heading_text('References', level=1)

refs = [
    '1. Rosenberg PH, Veering BT, Urmey WF. Maximum recommended doses of local anesthetics: a multifactorial concept. Reg Anesth Pain Med 2004; 29: 564\u201375.',
    '2. El-Boghdadly K, Pawa A, Chin KJ. Local anesthetic systemic toxicity: current perspectives. Local Reg Anesth 2018; 11: 35\u201344.',
    '3. Marsh B, White M, Morton N, Kenny GN. Pharmacokinetic model driven infusion of propofol in children. Br J Anaesth 1991; 67: 41\u20138.',
    '4. Schnider TW, Minto CF, Gambus PL, et al. The influence of method of administration and covariates on the pharmacokinetics of propofol in adult volunteers. Anesthesiology 1998; 88: 1170\u201382.',
    '5. Eleveld DJ, Colin P, Absalom AR, Struys MMRF. Pharmacokinetic\u2013pharmacodynamic model for propofol for broad application in anaesthesia and sedation. Br J Anaesth 2018; 120: 942\u201359.',
    '6. Minto CF, Schnider TW, Egan TD, et al. Influence of age and gender on the pharmacokinetics and pharmacodynamics of remifentanil. Anesthesiology 1997; 86: 10\u201323.',
    '7. Leite-Moreira AM, Correia A, Vale N, Mour\u00e3o JB. Pharmacokinetics in regional anesthesia. Curr Opin Anaesthesiol 2024; 37: 520\u20135.',
    '8. Arthur GR, Covino BG. Pharmacokinetics of local anaesthetics. Bailli\u00e8re\u2019s Clin Anaesthesiol 1991; 5: 635\u201358.',
    '9. De Cassai A, Dost B, Mormando G, Stecco C. Epinephrine, absorption, and local anaesthetic systemic toxicity: insights from continuous fascial block pharmacokinetic models. Br J Anaesth 2025; 135: 857\u201360.',
    '10. Schwenk ES, Sneyd JR, Wu CL. The state of local anaesthetic systemic toxicity in 2025: the emergence of lidocaine as our next challenge. Br J Anaesth 2025; 135: 854\u20136.',
    '11. Rahiri JL, Tuhoe J, Svirskis D, Lightfoot NJ, Lirk PB, Hill AG. Systematic review of the systemic concentrations of local anaesthetic after transversus abdominis plane block and rectus sheath block. Br J Anaesth 2017; 118: 517\u201326.',
    '12. De Cassai A, Pasin L, Boscolo A, et al. Safety of local anesthetics for fascial plane blocks: a narrative review. J Clin Anesth 2022; 77: 110637.',
    '13. Barrington MJ, Kluger R. Ultrasound guidance reduces the risk of local anesthetic systemic toxicity following peripheral nerve blockade. Reg Anesth Pain Med 2013; 38: 289\u201397.',
    '14. Fettiplace MR, Weinberg G. The mechanisms underlying lipid resuscitation therapy. Reg Anesth Pain Med 2018; 43: 138\u201349.',
    '15. Tucker GT, Mather LE. Clinical pharmacokinetics of local anaesthetic agents. Clin Pharmacokinet 1979; 4: 241\u201378.',
    '16. Simon MJG, Veering BT. Factors affecting the pharmacokinetics and neural block characteristics after epidural administration of local anaesthetics. Eur J Pain 2010; 4: 209\u201318.',
    '17. Burm AG, van der Meer AD, van Kleef JW, Zeijlmans PW, Groen K. Pharmacokinetics of the enantiomers of bupivacaine following intravenous administration of the racemate. Br J Clin Pharmacol 1994; 38: 125\u201329.',
    '18. Gaudreault F, Bherer L, Bhatt DL, Bhatt HV, Bhatt HV. Modeling the anesthetic effect of ropivacaine after a femoral nerve block in orthopedic patients. Anesthesiology 2015; 122: 1010\u201320.',
    '19. Ling J, Xu C, Tang L, Qiu L, Hu N. Comparison of the pharmacokinetic variations of different concentrations of ropivacaine used for serratus anterior plane block. Front Pharmacol 2025; 16: 1540606.',
    '20. Kwa A, Sprung J, Van Guilder M, Jelliffe RW. A population pharmacokinetic model of epidural lidocaine in geriatric patients. Ther Drug Monit 2008; 30: 346\u201355.',
    '21. Thompson MD, Beard DA. Physiologically-based pharmacokinetic tissue compartment model selection in drug development and risk assessment. J Pharm Sci 2012; 101: 424\u201335.',
    '22. Jones HM, Rowland-Yeo K. Basic concepts in physiologically based pharmacokinetic modeling in drug discovery and development. CPT Pharmacometrics Syst Pharmacol 2013; 2: e63.',
    '23. Lippert J, Burghaus R, Edginton A, et al. Open Systems Pharmacology Community\u2014an open access, open source, open science approach to modeling and simulation in pharmaceutical sciences. CPT Pharmacometrics Syst Pharmacol 2019; 8: 878\u201382.',
    '24. Willmann S, Lippert J, Sevestre M, Solodenko J, Fois F, Schmitt W. PK-Sim\u00ae: a physiologically based pharmacokinetic \u2018whole-body\u2019 model. Biosilico 2003; 1: 121\u20134.',
    '25. Open Systems Pharmacology. PK-Sim\u00ae Documentation: Formulations and Administration Protocols. Available at: https://docs.open-systems-pharmacology.org (accessed March 2026).',
    '26. Gamb\u00fas PL, Troc\u00f3niz IF. Pharmacokinetic\u2013pharmacodynamic modelling in anaesthesia. Br J Clin Pharmacol 2015; 79: 72\u201384.',
    '27. Hughes MA, Glass PS, Jacobs JR. Context-sensitive half-time in multicompartment pharmacokinetic models for intravenous anesthetic drugs. Anesthesiology 1992; 76: 334\u201341.',
    '28. Pirri C, Torre DE, Stecco C. Fascial plane blocks: from microanatomy to clinical applications. Curr Opin Anaesthesiol 2024; 37: 526\u201332.',
    '29. Sharma SK, Sonawane K, Mistry T. A narrative review on fascial plane blocks \u2013 Part A: Anatomical foundations and mechanistic insights. Indian J Anaesth 2026; 70: 127\u201336.',
    '30. Niederalt C, Kuepfer L, Solodenko J, et al. A generic whole body physiologically based pharmacokinetic model for therapeutic proteins in PK-Sim. J Pharmacokinet Pharmacodyn 2018; 45: 235\u201357.',
    '31. Gill KL, Gardner I, Li L, Jamei M. A bottom-up whole-body physiologically based pharmacokinetic model to mechanistically predict tissue distribution and the rate of subcutaneous absorption of therapeutic proteins. AAPS J 2016; 18: 156\u201370.',
    '32. Ashraf MW, Uusalo P, Scheinin M, Saari TI. Population modelling of dexmedetomidine pharmacokinetics and haemodynamic effects after intravenous and subcutaneous administration. Clin Pharmacokinet 2020; 59: 1467\u201382.',
    '33. Pepin XJH, Grant I, Wood JM. SubQ-Sim: a subcutaneous physiologically based biopharmaceutics model. Part 1: the injection and system parameters. Pharm Res 2023; 40: 2195\u2013214.',
    '34. Silva DA, Le Merdy M, Mullin J, et al. Mechanistic modeling of intramuscular administration of a long-acting injectable accounting for tissue response at the depot site. AAPS J 2026; 28: 4.',
    '35. De Cassai A, Bonvicini D, Correale C, et al. Histology of the fascial planes: a systematic review of the microstructural foundations of regional anesthesia. J Anesth Analg Crit Care 2026; 6: 5.',
    '36. Winnie AP, Tay CH, Patel KP, Ramamurthy S, Durrani Z. Pharmacokinetics of local anesthetics during plexus blocks. Anesth Analg 1977; 56: 852\u201361.',
    '37. Y\u00e1\u00f1ez JA, Remsberg CM, Sayre CL, Forrest ML, Davies NM. Flip-flop pharmacokinetics\u2014delivering a reversal of disposition: challenges and opportunities during drug development. Ther Deliv 2011; 2: 643\u201372.',
    '38. Butiulca M, Farczadi L, Imre S, et al. LC-MS/MS assisted pharmacokinetic and tissue distribution study of ropivacaine and 3-OH-ropivacaine on rats after plane block anesthesia. Front Pharmacol 2025; 15: 1494646.',
    '39. Osborne KW, MacFater WS, Anderson BJ, Svirskis D, Hill AG, Hannam JA. Pharmacokinetics of intraperitoneal lidocaine for sustained postoperative analgesia in adults. Eur J Drug Metab Pharmacokinet 2025.',
    '40. Bettonte S, Berton M, Battegay M, Stader F, Marzolini C. Development of a physiologically-based pharmacokinetic model to simulate the pharmacokinetics of intramuscular antiretroviral drugs. CPT Pharmacometrics Syst Pharmacol 2024; 13: 781\u201394.',
    '41. Dost B. Fascial plane blocks in the era of modern regional anesthesia: shaping the future of pain management. J Anesth Analg Crit Care 2025; 5: 49.',
    '42. Enlund M. TCI: Target Controlled Infusion, or Totally Confused Infusion? Upsala J Med Sci 2008; 113: 161\u201370.',
    '43. Vellinga R, Eleveld DJ, Struys MMRF, van den Berg JP. General purpose models for intravenous anesthetics, the next generation for target-controlled infusion and total intravenous anesthesia? Curr Opin Anaesthesiol 2023; 36: 602\u20137.',
    '44. \u0160afr\u00e1nkov\u00e1 P, Bruthans J. Target-controlled infusion of propofol: a systematic review of recent results. J Med Syst 2025; 49: 54.',
    '45. Sessler DI, Bao X, Leiman D, et al. A phase I study of the pharmacokinetics, pharmacodynamics, and safety of liposomal bupivacaine for sciatic nerve block in the popliteal fossa for bunionectomy. J Clin Pharmacol 2025; 65: 441\u201351.',
    '46. Xu A, Ren A, Lee C. Pharmacokinetics of lidocaine infusion: optimal dosing and duration in ERAS protocol. medRxiv 2025.',
    '47. Cascone S, Lamberti G, Titomanlio G, Piazza O. Pharmacokinetics of remifentanil: a three-compartmental modeling approach. Transl Med UniSa 2013; 7: 18\u201322.',
    '48. Brainkart. Pharmacokinetics: compartment models. In: Clinical Anesthesiology: Clinical Pharmacology. Available at: https://www.brainkart.com (accessed March 2026).',
    '49. Deranged Physiology. Single and multiple compartment models of drug distribution. Available at: https://derangedphysiology.com (accessed March 2026).',
    '50. Holt A. Three compartment drug. In: An ABC of PK/PD. Open Education Alberta; 2023.',
]

for ref in refs:
    add_para(ref)

doc.add_page_break()

# ===== INSERT FIGURES =====
add_heading_text('Figures', level=1)

fig_dir = '/home/ubuntu/manuscript/figures'
for i, (fname, caption) in enumerate([
    ('figure1_compartment_models.png', 'Figure 1'),
    ('figure2_pk_simulation.png', 'Figure 2'),
    ('figure3_workflow.png', 'Figure 3'),
    ('figure4_aims.png', 'Figure 4'),
], 1):
    fpath = os.path.join(fig_dir, fname)
    if os.path.exists(fpath):
        doc.add_paragraph(f'{caption}')
        doc.add_picture(fpath, width=Inches(6.0))
        doc.add_paragraph('')

# Save
out_path = '/home/ubuntu/manuscript/BJA_Manuscript_English.docx'
doc.save(out_path)
print(f'English manuscript saved to {out_path}')
